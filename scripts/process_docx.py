#!/usr/bin/env python3
"""
Procesador de documentos Word (.docx)
- Convierte comillas inglesas direccionales a comillas latinas
- Inserta notas al pie dentro del texto
"""

import sys
import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def convert_quotes_to_latina(text):
    """Convierte comillas inglesas direccionales a comillas latinas."""
    # Comillas inglesas direccionales: " (U+201C) y " (U+201D)
    # Comillas rectas: " (U+0022)
    # Comillas latinas: « (U+00AB) y » (U+00BB)
    
    original = text
    count = 0
    
    # Patrón para comillas inglesas direccionales: "texto"
    # Usamos un enfoque de pares: la primera comilla de apertura se convierte en «
    # y la siguiente de cierre se convierte en »
    
    result = []
    in_quotes = False
    i = 0
    while i < len(text):
        char = text[i]
        
        # Comilla inglesa de apertura "
        if char == '\u201C':  # "
            if not in_quotes:
                result.append('\u00AB')  # «
                in_quotes = True
                count += 1
            else:
                # Si ya estamos en comillas, cerramos y abrimos de nuevo
                result.append('\u00BB')  # »
                result.append('\u00AB')  # «
                count += 2
        # Comilla inglesa de cierre "
        elif char == '\u201D':  # "
            if in_quotes:
                result.append('\u00BB')  # »
                in_quotes = False
                count += 1
            else:
                # Si no estamos en comillas, la dejamos como está
                result.append(char)
        # Comillas rectas " - intentar convertir si delimitan texto
        elif char == '\u0022':  # "
            if not in_quotes:
                result.append('\u00AB')  # «
                in_quotes = True
                count += 1
            else:
                result.append('\u00BB')  # »
                in_quotes = False
                count += 1
        else:
            result.append(char)
        i += 1
    
    return ''.join(result), count


def get_footnotes(doc):
    """Extrae las notas al pie del documento."""
    footnotes = {}
    
    # Buscar el elemento de notas al pie en el documento
    try:
        # Las notas al pie están en doc.part.rels
        for rel in doc.part.rels.values():
            if "footnotes" in rel.reltype:
                # Cargar el XML de notas al pie
                footnotes_part = rel.target_part
                footnotes_xml = footnotes_part._blob
                
                # Parsear el XML
                from lxml import etree
                root = etree.fromstring(footnotes_xml)
                
                # Extraer cada nota al pie
                for footnote in root.findall('.//' + qn('w:footnote')):
                    note_id = footnote.get(qn('w:id'))
                    if note_id and note_id not in ['-1', '0']:  # Ignorar separadores
                        # Extraer el texto de la nota
                        text_parts = []
                        for t in footnote.findall('.//' + qn('w:t')):
                            if t.text:
                                text_parts.append(t.text)
                        if text_parts:
                            footnotes[int(note_id)] = ' '.join(text_parts)
    except Exception as e:
        print(f"Debug: Error extrayendo notas al pie: {e}", file=sys.stderr)
    
    return footnotes


def process_paragraph(paragraph, footnotes):
    """Procesa un párrafo: convierte comillas e inserta notas al pie."""
    footnotes_inserted = 0
    quotes_converted = 0
    
    # Verificar si el párrafo tiene referencias a notas al pie
    footnote_refs = []
    
    # Buscar referencias a notas al pie en el XML del párrafo
    for run in paragraph.runs:
        for elem in run._element.iter():
            if elem.tag == qn('w:footnoteReference'):
                note_id = elem.get(qn('w:id'))
                if note_id:
                    footnote_refs.append((run, int(note_id)))
    
    # Procesar cada run del párrafo
    for run in paragraph.runs:
        if run.text:
            converted_text, count = convert_quotes_to_latina(run.text)
            if count > 0:
                run.text = converted_text
                quotes_converted += count
    
    # Insertar notas al pie después del último run que tenga referencia
    if footnote_refs and footnotes:
        # Procesar de atrás hacia adelante para no afectar los índices
        footnote_refs.reverse()
        
        for run, note_id in footnote_refs:
            if note_id in footnotes:
                note_text = footnotes[note_id]
                # Añadir el texto de la nota al final del run
                if run.text:
                    run.text = run.text.rstrip() + f' [Nota al pie: {note_text}]'
                else:
                    run.text = f' [Nota al pie: {note_text}]'
                footnotes_inserted += 1
    
    return quotes_converted, footnotes_inserted


def process_document(input_path, output_path):
    """Procesa un documento Word completo."""
    try:
        # Cargar el documento
        doc = Document(input_path)
        
        # Extraer notas al pie
        footnotes = get_footnotes(doc)
        print(f"Debug: Encontradas {len(footnotes)} notas al pie", file=sys.stderr)
        
        total_quotes = 0
        total_footnotes = 0
        
        # Procesar cada párrafo
        for paragraph in doc.paragraphs:
            q, f = process_paragraph(paragraph, footnotes)
            total_quotes += q
            total_footnotes += f
        
        # También procesar tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        q, f = process_paragraph(paragraph, footnotes)
                        total_quotes += q
                        total_footnotes += f
        
        # Guardar el documento procesado
        doc.save(output_path)
        
        return {
            'success': True,
            'quotesConverted': total_quotes,
            'footnotesInserted': total_footnotes
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Uso: python process_docx.py <input_file> <output_file>", file=sys.stderr)
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    result = process_document(input_file, output_file)
    
    if result['success']:
        print(f"SUCCESS|{result['quotesConverted']}|{result['footnotesInserted']}")
    else:
        print(f"ERROR|{result['error']}")
